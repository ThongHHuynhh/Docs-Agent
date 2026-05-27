import io
import os
import yaml
import subprocess
from pathlib import Path
from docx import Document
from google import genai
import re

client = genai.Client(api_key=os.environ["GEMINI_API_KEY"])

ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT.parent / "V3" /"Docs"
TEMPLATES_DIR = ROOT  / "templates"
CUSTOMERS_DIR = ROOT / "customers"
OUTPUTS_DIR = ROOT / "outputs"

IMAGE_PLACEHOLDER_REGEX = re.compile(r"^\[\[IMAGE:(.+?)\]\]$")


def extract_docx_elements(file_path):
    doc = Document(file_path)

    def has_tag(element, local_name):
        return any(child.tag.endswith('}' + local_name) for child in element.iter())

    def find_blip_rel_id(element):
        rel_attr = '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed'
        for child in element.iter():
            if child.tag.endswith('}blip'):
                return child.get(rel_attr)
        return None

    elements = []

    for paragraph in doc.paragraphs:
        text_parts = []
        for run in paragraph.runs:
            if has_tag(run.element, 'pic'):
                if text_parts:
                    elements.append({"type": "text", "text": "".join(text_parts)})
                    text_parts = []

                rel_id = find_blip_rel_id(run.element)
                if rel_id and rel_id in doc.part.related_parts:
                    image_part = doc.part.related_parts[rel_id]
                    elements.append({
                        "type": "image",
                        "image_name": Path(image_part.partname).name,
                        "image_bytes": image_part.blob,
                        "source": file_path.name,
                    })
            else:
                text_parts.append(run.text or "")

        if text_parts:
            elements.append({"type": "text", "text": "".join(text_parts)})

        # Preserve paragraph breaks as their own element so placeholders stay aligned.
        elements.append({"type": "text", "text": "\n"})

    return elements


def read_universal_docs():
    content = []
    image_map = {}

    supported_ext = [".md", ".txt", ".docx"]
    #supported_ext = [".docx"]

    for file_path in DOCS_DIR.rglob("*"):
        if file_path.is_file() and file_path.suffix.lower() in supported_ext:
            if file_path.suffix.lower() in [".md", ".txt"]:
                text = file_path.read_text(encoding="utf-8", errors="ignore")
                content.append(f"\n\n--- SOURCE FILE: {file_path.name} ---\n{text}")
            elif file_path.suffix.lower() == ".docx":
                elements = extract_docx_elements(file_path)
                block_lines = [f"--- SOURCE FILE: {file_path.name} ---"]

                for index, element in enumerate(elements):
                    if element["type"] == "text":
                        block_lines.append(element["text"])
                    else:
                        marker = f"[[IMAGE:{file_path.name}:{index}:{element['image_name']}]]"
                        block_lines.append(marker)
                        image_map[marker] = element

                content.append("\n" + "\n".join(block_lines))

    print(f"Loaded {len(content)} universal documentation files.")
    return "\n".join(content), image_map

# Load customer configuration from YAML file
def load_customer(customer_file):
    with open(customer_file, "r", encoding="utf-8") as f:
        return yaml.safe_load(f)


def generate_custom_doc(universal_docs, customer):
    prompt = f"""
You are a technical documentation agent.

You must customize product documentation using ONLY the universal documentation below. 
The given documentation is customized for each customer, based on that make a general version of the documentation.

Universal documentation:
{universal_docs}

Customer configuration:
{customer}


Task:
Generalize the given documents and create a customized customer-use product document.

Rules:
- Use clear headings.
- Be professional.
- Do not invent unsupported features.
- Mention the customer name where appropriate.
- Be simple and straightforward, with customer perspective.
- Preserve image placeholders of the form [[IMAGE:<source-file>:<index>:<name>]].
- Do not modify, remove, or rewrite those placeholders.
- Keep each image placeholder on its own line so the final document can place the original image there.
"""

    response = client.models.generate_content(
        model="gemini-2.5-flash",
        contents=prompt
    )

    return response.text

def add_bold_text(doc, text):
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = doc.add_run(part[2:-2])
            run.bold = True
        else:
            doc.add_run(part)

def save_docx(content, output_path, customer_name, image_map=None):
    doc = Document()
    doc.add_heading(f"{customer_name} - Customized Product Documentation", level=0).bold = True

    for line in content.split("\n"):
        line = line.strip()

        if not line:
            continue

        image_element = image_map.get(line) if image_map else None
        if image_element:
            image_stream = io.BytesIO(image_element["image_bytes"])
            doc.add_picture(image_stream)
            continue

        if line.startswith("# "):
            doc.add_heading(line.replace("# ", ""), level=1)
        elif line.startswith("## "):
            doc.add_heading(line.replace("## ", ""), level=2)
        elif line.startswith("### "):
            doc.add_heading(line.replace("### ", ""), level=3)
        elif line.startswith("* "):
            bullet_text = line[2:].strip()
            paragraph = doc.add_paragraph(style="List Bullet")
            add_bold_text(paragraph, bullet_text)
        else:
            p = doc.add_paragraph()
            add_bold_text(p, line)

    doc.save(output_path)


def convert_to_pdf(docx_path, output_dir):
    try:
        subprocess.run([
            "soffice",
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(output_dir),
            str(docx_path)
        ], check=True)
    except FileNotFoundError:
        print("PDF conversion skipped: LibreOffice is not installed or 'soffice' is not in PATH.")


def generate_for_customer(customer_file):
    customer = load_customer(customer_file)
    universal_docs, image_map = read_universal_docs()

    customer_name = customer["customer_name"]
    #replace spaces with underscores for output file naming, and convert to lowercase
    output_name = customer.get("output_name", customer_name.lower().replace(" ", "_"))

    output_folder = OUTPUTS_DIR / output_name
    output_folder.mkdir(parents=True, exist_ok=True)

    content = generate_custom_doc(universal_docs, customer)

    docx_path = output_folder / f"{output_name}.docx"
    save_docx(content, docx_path, customer_name, image_map)

    convert_to_pdf(docx_path, output_folder)

    print(f"Generated files for {customer_name}")
    print(f"DOCX: {docx_path}")


def main():
    customer_files = list(CUSTOMERS_DIR.glob("*.yaml"))

    if not customer_files:
        print("No customer YAML files found.")
        return

    for customer_file in customer_files:
        generate_for_customer(customer_file)

if __name__ == "__main__":
    main()